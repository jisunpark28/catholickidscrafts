#!/usr/bin/env python3
"""Generate one-page customer support requests document with graph."""

import os
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "assignments")
GRAPH_PATH = os.path.join(OUTPUT_DIR, "support-requests-graph.png")
DOC_PATH = os.path.join(OUTPUT_DIR, "Customer-Support-Requests-Jan-Jun-2026.docx")

months = ["January", "February", "March", "April", "May", "June"]
requests = [120, 136, 151, 147, 169, 188]

os.makedirs(OUTPUT_DIR, exist_ok=True)

# Create column graph
fig, ax = plt.subplots(figsize=(8, 4.5))
bars = ax.bar(months, requests, color="#2E5EAA", edgecolor="#1a3d6e", linewidth=0.8, width=0.65)

for bar, value in zip(bars, requests):
    ax.text(
        bar.get_x() + bar.get_width() / 2,
        bar.get_height() + 2,
        str(value),
        ha="center",
        va="bottom",
        fontsize=11,
        fontweight="bold",
    )

ax.set_title(
    "Customer-Support Requests, January to June 2026",
    fontsize=14,
    fontweight="bold",
    pad=14,
)
ax.set_xlabel("Month", fontsize=12, labelpad=8)
ax.set_ylabel("Number of Requests", fontsize=12, labelpad=8)
ax.set_ylim(0, 210)
ax.tick_params(axis="both", labelsize=11)
ax.spines["top"].set_visible(False)
ax.spines["right"].set_visible(False)
ax.grid(axis="y", linestyle="--", alpha=0.35)
plt.xticks(rotation=0)
plt.tight_layout()
plt.savefig(GRAPH_PATH, dpi=200, bbox_inches="tight", facecolor="white")
plt.close()

explanation = (
    "Figure 1 shows customer-support requests from January to June 2026. Overall, "
    "requests followed an upward trend, rising from 120 in January to 188 in June. "
    "The only decrease occurred in April, when requests dropped from 151 in March to "
    "147. After that brief dip, volume climbed again in May and June. June recorded "
    "the highest number of requests at 188. Together, these figures suggest steadily "
    "growing demand for customer support during the first half of the year."
)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(12)

doc.add_picture(GRAPH_PATH, width=Inches(6.5))

caption = doc.add_paragraph()
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = caption.add_run(
    "Figure 1: Customer-Support Requests, January to June 2026"
)
run.bold = True
run.font.size = Pt(12)

source = doc.add_paragraph()
source.alignment = WD_ALIGN_PARAGRAPH.CENTER
source_run = source.add_run("Source: Course-provided data.")
source_run.italic = True
source_run.font.size = Pt(12)

doc.add_paragraph()

para = doc.add_paragraph(explanation)
para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
for run in para.runs:
    run.font.size = Pt(12)

doc.save(DOC_PATH)
print(f"Created: {DOC_PATH}")
print(f"Graph: {GRAPH_PATH}")
print(f"Explanation word count: {len(explanation.split())}")
